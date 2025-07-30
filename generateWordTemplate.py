#!/usr/bin/env python3
"""
Modify an existing Word document to add docxtpl template placeholders
This opens an existing SOP template and adds Jinja2 placeholders
"""

from docxtpl import DocxTemplate
from docx import Document
from pathlib import Path
import re


def add_placeholders_to_existing_doc(input_path, output_path):
    """
    Open an existing Word document and add template placeholders

    Args:
        input_path (str): Path to existing Word document
        output_path (str): Path where modified template will be saved
    """

    # Open the existing document with python-docx first to modify content
    doc = Document(input_path)

    print(f"📖 Opening existing document: {input_path}")
    print(f"Found {len(doc.paragraphs)} paragraphs")

    # Dictionary of text to replace with placeholders
    replacements = {
        # You can customize these based on your actual document content
        'OBJECTIVE_PLACEHOLDER': '{{objective}}',
        'SCOPE_PLACEHOLDER': '{{scope}}',
        'RESPONSIBILITIES_PLACEHOLDER': '{{responsibilities}}',
        'DEFINITIONS_PLACEHOLDER': '{{definitions}}',
        'MANGO_FILE_ID': '{{mango_pptx_id}}',
        'PROCEDURE_TITLE': '{{pptx_title}}',
        'DATE_PLACEHOLDER': '{{date}}',
        'USER_NAME_PLACEHOLDER': '{{user_name}}'
    }

    # Process each paragraph
    for i, paragraph in enumerate(doc.paragraphs):
        original_text = paragraph.text
        modified = False

        # Check for text that should be replaced with placeholders
        for placeholder_text, jinja_placeholder in replacements.items():
            if placeholder_text in original_text:
                paragraph.text = original_text.replace(placeholder_text, jinja_placeholder)
                modified = True
                print(f"  ✏️  Paragraph {i}: Replaced '{placeholder_text}' with '{jinja_placeholder}'")

        # You can also add more sophisticated pattern matching
        # For example, if you want to replace specific patterns:
        if 'See Mango File' in original_text and '{{mango_pptx_id}}' not in original_text:
            # Replace a pattern like "See Mango File XYZ-123 -- Title Here"
            new_text = re.sub(
                r'See Mango File [A-Z0-9-]+ -- .*?\.',
                'See Mango File {{mango_pptx_id}} -- {{pptx_title}}.',
                original_text
            )
            if new_text != original_text:
                paragraph.text = new_text
                modified = True
                print(f"  ✏️  Paragraph {i}: Updated Mango File reference")

    # Process tables (for revision sheet)
    for table_idx, table in enumerate(doc.tables):
        print(f"📋 Processing table {table_idx + 1} ({len(table.rows)} rows, {len(table.columns)} columns)")

        # If this looks like a revision table, add placeholders
        if len(table.columns) == 4 and len(table.rows) > 1:
            # Check if first row has revision table headers
            first_row_text = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if any(word in ' '.join(first_row_text) for word in ['date', 'version', 'changes', 'changed']):
                print("  📝 Found revision table, adding placeholders...")

                # Add template row (usually row 1, after headers)
                if len(table.rows) > 1:
                    template_row = table.rows[1]
                    template_row.cells[0].text = '{{date}}'
                    template_row.cells[1].text = '1.0'
                    template_row.cells[2].text = 'Initial Release'
                    template_row.cells[3].text = '{{user_name}}'
                    print("  ✅ Added revision table placeholders")

    # Save the modified document
    doc.save(output_path)
    print(f"💾 Saved modified document to: {output_path}")

    # Now convert to DocxTemplate for testing
    template = DocxTemplate(output_path)

    # Test with sample data
    test_context = {
        'objective': 'This SOP describes the procedure for testing template generation.',
        'scope': 'This SOP applies to all Cassette Manufacturing operations at the US Stoneham facility.',
        'responsibilities': 'Technician: Performs the procedure according to this SOP\nEngineer: Reviews and approves results',
        'definitions': 'N/A - No special definitions required for this procedure.',
        'mango_pptx_id': 'MANGO-2024-001',
        'pptx_title': 'Cassette Assembly Test Procedure',
        'date': '12/15/2024',
        'user_name': 'Manufacturing Engineer'
    }

    # Create test output
    test_output_path = output_path.replace('.docx', '_TEST_OUTPUT.docx')
    template.render(test_context)
    template.save(test_output_path)

    print(f"🎯 Created test output: {test_output_path}")
    return template


def interactive_placeholder_replacement(input_path, output_path):
    """
    Interactive version that lets you see the content and decide what to replace
    """
    doc = Document(input_path)

    print(f"\n📖 Analyzing document: {input_path}")
    print("=" * 60)

    # Show document structure
    print("DOCUMENT STRUCTURE:")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():  # Only show non-empty paragraphs
            print(f"  [{i:2d}] {paragraph.text[:80]}{'...' if len(paragraph.text) > 80 else ''}")

    print(f"\nTABLES FOUND: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"  Table {i + 1}: {len(table.rows)} rows × {len(table.columns)} columns")
        if table.rows:
            first_row = ' | '.join([cell.text[:20] for cell in table.rows[0].cells])
            print(f"    Headers: {first_row}")

    print("\n" + "=" * 60)
    print("Now you can manually edit the document to add placeholders like:")
    print("  {{objective}}, {{scope}}, {{responsibilities}}, etc.")
    print("\nOr run the automatic replacement function.")

    # Ask user what they want to do
    choice = input("\nDo you want to (1) auto-replace or (2) manual edit? [1/2]: ")

    if choice == "1":
        return add_placeholders_to_existing_doc(input_path, output_path)
    else:
        print(f"\n📝 Please manually edit {input_path} to add placeholders, then save as {output_path}")
        print("Then you can use DocxTemplate to render it with data.")
        return None


def find_and_replace_specific_content(input_path, output_path, replacements_dict):
    """
    More targeted replacement based on exact text matches

    Args:
        input_path: Path to existing document
        output_path: Path for modified template
        replacements_dict: Dict of {original_text: placeholder} pairs
    """
    doc = Document(input_path)

    print(f"🔍 Searching for specific content to replace...")

    replacements_made = 0

    # Process paragraphs
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        new_text = original_text

        for original, placeholder in replacements_dict.items():
            if original in new_text:
                new_text = new_text.replace(original, placeholder)
                replacements_made += 1
                print(f"  ✅ Replaced: '{original}' → '{placeholder}'")

        if new_text != original_text:
            paragraph.text = new_text

    # Process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                new_text = original_text

                for original, placeholder in replacements_dict.items():
                    if original in new_text:
                        new_text = new_text.replace(original, placeholder)
                        replacements_made += 1
                        print(f"  ✅ Replaced in table: '{original}' → '{placeholder}'")

                if new_text != original_text:
                    cell.text = new_text

    doc.save(output_path)
    print(f"\n💾 Made {replacements_made} replacements and saved to: {output_path}")

    return DocxTemplate(output_path)


if __name__ == "__main__":
    print("SOP Template Modifier")
    print("=" * 50)

    # Define your paths
    input_document = "path/to/your/existing_sop.docx"  # Change this!
    output_template = "templates/LS_SOP_Template_Modified.docx"

    # Make sure templates directory exists
    Path("templates").mkdir(exist_ok=True)

    # Example 1: Automatic replacement with common patterns
    if Path(input_document).exists():
        print("Method 1: Automatic placeholder insertion")
        template = add_placeholders_to_existing_doc(input_document, output_template)

        print("\n" + "=" * 50)
        print("Method 2: Specific content replacement")

        # Define exactly what text to replace
        specific_replacements = {
            "Enter objective here": "{{objective}}",
            "Enter scope here": "{{scope}}",
            "List responsibilities here": "{{responsibilities}}",
            "Define terms here": "{{definitions}}",
            "John Doe": "{{user_name}}",
            "01/01/2024": "{{date}}"
        }

        template2 = find_and_replace_specific_content(
            input_document,
            output_template.replace('.docx', '_specific.docx'),
            specific_replacements
        )

    else:
        print(f"❌ Input document not found: {input_document}")
        print("Please update the 'input_document' path to point to your existing SOP template.")
        print("\nExample usage:")
        print("1. Save your existing Word SOP as 'my_sop_template.docx'")
        print("2. Update input_document = 'my_sop_template.docx'")
        print("3. Run this script")
        print("4. The script will add {{placeholders}} where needed")
        print("5. Use the resulting template with docxtpl in your web app")
#!/usr/bin/env python3
"""
Analyze how docxtpl reads and parses Word document templates
This script examines a Word document and shows all template variables found
"""

from docxtpl import DocxTemplate
from docx import Document
from pathlib import Path
import re
import os
import json
from typing import Dict, List, Any
import xml.etree.ElementTree as ET


def analyze_docxtpl_template(template_path: str):
    """
    Comprehensive analysis of how docxtpl reads a Word template

    Args:
        template_path: Path to the Word document template
    """

    print("=" * 80)
    print(f"DOCXTPL TEMPLATE ANALYSIS: {template_path}")
    print("=" * 80)

    # Check if file exists
    if not Path(template_path).exists():
        print(f"❌ File not found: {template_path}")
        return

    try:
        # 1. Load with DocxTemplate
        template = DocxTemplate(template_path)
        print("✅ Successfully loaded template with DocxTemplate")

        # 2. Load with python-docx for comparison
        doc = Document(template_path)
        print("✅ Successfully loaded document with python-docx")

    except Exception as e:
        print(f"❌ Error loading template: {e}")
        return

    print("\n" + "=" * 60)
    print("1. BASIC DOCUMENT STRUCTURE")
    print("=" * 60)

    print(f"📄 Paragraphs: {len(doc.paragraphs)}")
    print(f"📋 Tables: {len(doc.tables)}")
    print(f"🖼️  Inline shapes: {len(doc.inline_shapes) if hasattr(doc, 'inline_shapes') else 'N/A'}")

    print("\n" + "=" * 60)
    print("2. TEMPLATE VARIABLES FOUND")
    print("=" * 60)

    # Extract all Jinja2 template variables using regex
    all_text = extract_all_text_from_doc(doc)
    variables = find_jinja_variables(all_text)

    if variables:
        print(f"🎯 Found {len(variables)} unique template variables:")
        for var_type, var_list in variables.items():
            if var_list:
                print(f"\n  {var_type.upper()}:")
                for var in sorted(set(var_list)):
                    print(f"    • {var}")
    else:
        print("❌ No Jinja2 template variables found")

    print("\n" + "=" * 60)
    print("3. PARAGRAPH-BY-PARAGRAPH ANALYSIS")
    print("=" * 60)

    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:  # Only show non-empty paragraphs
            jinja_vars = find_jinja_variables(text)
            has_templates = any(jinja_vars.values())

            status = "🎯" if has_templates else "📝"
            print(f"{status} Paragraph {i:2d}: {text[:60]}{'...' if len(text) > 60 else ''}")

            if has_templates:
                for var_type, vars_found in jinja_vars.items():
                    if vars_found:
                        print(f"    └─ {var_type}: {vars_found}")

    print("\n" + "=" * 60)
    print("4. TABLE ANALYSIS")
    print("=" * 60)

    for table_idx, table in enumerate(doc.tables):
        print(f"📋 Table {table_idx + 1}: {len(table.rows)} rows × {len(table.columns)} columns")

        has_template_vars = False
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    jinja_vars = find_jinja_variables(cell_text)
                    if any(jinja_vars.values()):
                        has_template_vars = True
                        print(f"  🎯 Cell ({row_idx}, {col_idx}): {cell_text}")
                        for var_type, vars_found in jinja_vars.items():
                            if vars_found:
                                print(f"      └─ {var_type}: {vars_found}")

        if not has_template_vars:
            # Show first row as headers
            if table.rows:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                print(f"    Headers: {' | '.join(headers)}")

    print("\n" + "=" * 60)
    print("5. TEST TEMPLATE RENDERING")
    print("=" * 60)

    # Try to render with dummy data
    test_context = create_test_context(variables)

    if test_context:
        print("🧪 Testing template rendering with dummy data:")
        print(f"   Context variables: {list(test_context.keys())}")

        try:
            # Create a copy for testing
            test_template = DocxTemplate(template_path)
            test_template.render(test_context)

            # Save to memory stream to test if it works
            from io import BytesIO
            output_stream = BytesIO()
            test_template.save(output_stream)

            print("✅ Template rendering successful!")
            print(f"   Generated document size: {len(output_stream.getvalue())} bytes")

        except Exception as e:
            print(f"❌ Template rendering failed: {e}")
            print("   This might indicate missing variables or syntax errors")

    print("\n" + "=" * 60)
    print("6. DOCXTPL INTERNAL ANALYSIS")
    print("=" * 60)

    # Try to access internal template structure
    try:
        # Look at the internal Jinja2 environment
        if hasattr(template, 'docx'):
            print("📦 DocxTemplate internal structure:")
            print(f"   • Document object: {type(template.docx)}")

        # Check for Jinja2 environment
        if hasattr(template, 'jinja_env'):
            print(f"   • Jinja2 environment: {type(template.jinja_env)}")

    except Exception as e:
        print(f"⚠️  Could not analyze internal structure: {e}")

    return {
        'variables': variables,
        'paragraph_count': len(doc.paragraphs),
        'table_count': len(doc.tables),
        'test_context': test_context
    }


def extract_all_text_from_doc(doc: Document) -> str:
    """Extract all text from document including paragraphs and tables"""
    all_text = []

    # Get text from paragraphs
    for paragraph in doc.paragraphs:
        all_text.append(paragraph.text)

    # Get text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)

    return '\n'.join(all_text)


def find_jinja_variables(text: str) -> Dict[str, List[str]]:
    """
    Find all Jinja2 template variables in text
    Returns categorized variables by type
    """

    # Regular expressions for different Jinja2 constructs
    patterns = {
        'variables': r'\{\{\s*([^}]+)\s*\}\}',  # {{ variable }}
        'blocks': r'\{\%\s*([^%]+)\s*\%\}',  # {% block %}
        'comments': r'\{\#\s*([^#]+)\s*\#\}',  # {# comment #}
        'filters': r'\{\{\s*([^|]+)\|([^}]+)\s*\}\}',  # {{ var|filter }}
    }

    results = {}

    for pattern_type, pattern in patterns.items():
        matches = re.findall(pattern, text)

        if pattern_type == 'filters':
            # For filters, we get tuples (variable, filter)
            variables = [match[0].strip() for match in matches]
            filters = [match[1].strip() for match in matches]
            results[pattern_type] = [f"{var}|{filt}" for var, filt in zip(variables, filters)]
        else:
            results[pattern_type] = [match.strip() for match in matches]

    return results


def create_test_context(variables: Dict[str, List[str]]) -> Dict[str, Any]:
    """Create dummy context data for testing template rendering"""

    context = {}

    # Extract simple variables (not blocks or comments)
    simple_vars = variables.get('variables', [])

    for var in simple_vars:
        # Clean up variable name (remove filters, etc.)
        clean_var = var.split('|')[0].split('.')[0].strip()

        # Generate appropriate dummy data based on variable name
        if any(keyword in clean_var.lower() for keyword in ['date', 'time']):
            context[clean_var] = '2024-12-15'
        elif any(keyword in clean_var.lower() for keyword in ['name', 'user', 'author']):
            context[clean_var] = 'Test User'
        elif any(keyword in clean_var.lower() for keyword in ['id', 'number', 'version']):
            context[clean_var] = 'TEST-001'
        elif any(keyword in clean_var.lower() for keyword in ['title', 'subject']):
            context[clean_var] = 'Test Document Title'
        else:
            context[clean_var] = f'Test value for {clean_var}'

    return context


def compare_extraction_methods(template_path: str):
    """
    Compare different methods of extracting variables from templates
    """

    print("\n" + "=" * 60)
    print("7. COMPARISON OF EXTRACTION METHODS")
    print("=" * 60)

    # Method 1: Manual regex on raw text
    doc = Document(template_path)
    raw_text = extract_all_text_from_doc(doc)
    manual_vars = find_jinja_variables(raw_text)

    print("📝 Method 1: Manual regex extraction")
    print(f"   Variables found: {len(manual_vars.get('variables', []))}")
    print(f"   Variables: {manual_vars.get('variables', [])}")

    # Method 2: Try to access docxtpl internals (if possible)
    try:
        template = DocxTemplate(template_path)
        print("\n🔧 Method 2: DocxTemplate internal analysis")
        print("   (Limited access to internals in current docxtpl version)")

    except Exception as e:
        print(f"\n❌ Method 2 failed: {e}")

    return manual_vars


def analyze_template_file(file_path: str):
    """Main analysis function"""

    if not Path(file_path).exists():
        print(f"Creating a sample template for testing...")

    # Run the analysis
    results = analyze_docxtpl_template(file_path)

    # Additional comparison
    compare_extraction_methods(file_path)

    # Summary
    print("\n" + "=" * 80)
    print("ANALYSIS SUMMARY")
    print("=" * 80)

    if results:
        total_vars = sum(len(vars_list) for vars_list in results['variables'].values())
        print(f"📊 Total template constructs found: {total_vars}")
        print(f"📄 Document structure: {results['paragraph_count']} paragraphs, {results['table_count']} tables")
        if results['test_context']:
            print(f"🧪 Test context created with {len(results['test_context'])} variables")

    return results

if __name__ == "__main__":
    print("DocxTemplate Variable Analyzer")
    print("=" * 50)

    # You can specify your template file here
    template_file = "LS Standard Operating Procedure Template 20107897 - v6.0 - 20107897.docx"  # Change this to your file

    # Alternative: Ask user for file path
    import sys

    if len(sys.argv) > 1:
        template_file = sys.argv[1]
    else:
        template_file = input("Enter absolute path to Word template (or press Enter for default): ").strip()
        if not template_file:
            template_file = os.path.abspath(template_file)

    # Run the analysis
    try:
        results = analyze_template_file(template_file)
        print(f"\n✅ Analysis complete for: {template_file}")

    except Exception as e:
        print(f"❌ Analysis failed: {e}")
        import traceback

        traceback.print_exc()